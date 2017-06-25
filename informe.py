#!/usr/bin/env python
# -*- coding: utf-8 -*

# https://python-docx.readthedocs.io/en/latest/
import plotly.offline as offline
from docx import Document
from docx.shared import Inches
import json


def decodificarFecha (str_fecha = "20190711"): # YYYYMMDD
	str_fecha = str(str_fecha)

	str_año = str_fecha[:4]
	str_mes = str_fecha[4:6]
	str_dia = str_fecha[6:8]

	dic_meses = {
		"01": "Enero", 
		"02": "Febrero",
		"03": "Marzo",
		"04": "Abril",
		"05": "Mayo",
		"06": "Junio",
		"07": "Julio",
		"08": "Agosto",
		"09": "Septiembre",
		"10": "Octubre",
		"11": "Noviembre",
		"12": "Diciembre"}


	return str_dia + " de " + dic_meses[str_mes] + " del " + str_año



def graficarPapelera (str_agenda = "agenda.json"):
	json_agenda     = {}
	list_fechas     = []
	list_papeleras  = [] 


	# Abrir el fichero con la agenda y extraer datos en formato json
	with open (str_agenda) as json_file:
		json_agenda = json.load(json_file)
	json_file.close()

	# Obtener fechas de la agenda y ordenar ascendentemente
	list_indices = sorted(list(json_agenda))

	# Obtener papeleras válidas y sus respectivas fechas
	for str_fecha in list_indices:
		try:
			int_papelera = str(json_agenda[str_fecha]["papelera"])
			if (int_papelera != "-1"):
				list_fechas.append("_"+str_fecha+"_")
				list_papeleras.append(int_papelera)
		except Exception as e:
			raise e
	
	list_papeleras.append("120")
	list_papeleras.append("0")
	# Dibujar la gráfica
	offline.plot({'data': [{'x': list_fechas, 'y': list_papeleras}], 
		            'layout': {
		            	'title': 'Nivel de la papelera ' + "(Junio 2017)", ### BORRA, INCOMPLETO, HAY QUE PONER LA FECHA DINAMICA 
		            	'xaxis': {
		            		"zeroline": True,
		                "showline": True
		            	},
		              'yaxis': {
		                "zeroline": True,
		                "showline": True,
		              	"range": [0, 120]}}},
		           auto_open=True, image = 'png', image_filename='papelera',
							 output_type='file', image_width=720, image_height=640)



def generarInforme (str_agenda = "agenda.json"):

	# Abrir el fichero con la agenda y extraer datos en formato json
	with open (str_agenda) as json_file:
		json_agenda = json.load(json_file)
	json_file.close()

	# Obtener fechas de la agenda y ordenar ascendentemente
	list_indices = sorted(list(json_agenda))

	doc = Document() # Crear documento vacío
	doc.add_heading('Informe de alergia \n' + 
	                decodificarFecha(list_indices[0]) + '  ---  ' + 
	                decodificarFecha(list_indices[-1]), 0)


	# Añadir registros al documento
	for str_fecha in list_indices:

		# Título de la entrada, fecha
		doc.add_heading(decodificarFecha(str_fecha), 1)

		# Sintomatología
		try:
			str_sintomas = str(json_agenda[str_fecha]["sintomas"])
			p = doc.add_paragraph()
			p.add_run("Síntomas: \n").bold = True
			p.add_run(str_sintomas).italic = True
		except Exception as e:
			pass


		# Medicación
		try:
			str_desayuno	= str(json_agenda[str_fecha]["medicacion"]["desayuno"])
			str_comida		= str(json_agenda[str_fecha]["medicacion"]["comida"])
			str_cena 			= str(json_agenda[str_fecha]["medicacion"]["cena"])
			
			if ((str_desayuno + str_comida + str_cena) != ""):
				p = doc.add_paragraph()
				p.add_run("Medicación:").bold = True
				if (str_desayuno != ""):
					doc.add_paragraph("desayuno: " + str_desayuno, style="ListBullet")

				if (str_comida != ""):
					doc.add_paragraph("comida: " + str_comida, style="ListBullet")

				if (str_cena != ""):
					doc.add_paragraph("cena: " + str_cena, style="ListBullet")

		except Exception as e:
			pass

	  
		# Estado de la papelera
		try:
			str_porcentaje = str(json_agenda[str_fecha]["papelera"])

			if (str_porcentaje != "-1"):
				p = doc.add_paragraph()
				p.add_run("Papelera: ").bold = True
				p.add_run(str_porcentaje + "%")
		except Exception as e:
			pass


		# Valoración
		try:
			dic_valoracion = {"1": "Bueno", "2": "Regular", "3": "Malo"}
			str_valoracion = str(json_agenda[str_fecha]["valoracion"])
			str_valoracion = dic_valoracion[str_valoracion]

			p = doc.add_paragraph()
			p.add_run("Valoración: ").bold = True
			p.add_run(str_valoracion)
		except Exception as e:
			pass

		#print(json_agenda[str_fecha])


	# Grafica de papelera
	try:
		doc.add_paragraph()
		doc.add_picture('papelera.png', width=Inches(7))
	except Exception as e:
		raise e

	# Guardar documento
	doc.save("informeAlergia.docx")





